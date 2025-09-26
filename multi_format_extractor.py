"""
KRA Data Extraction System - Enhanced Multi-Format Processor
===========================================================

Processes multiple documents (PDF, Word) from folders and extracts KRA data.
Supports both individual file uploads and batch folder processing.

Author: Groot
Date: September 20, 2025
"""


import streamlit as st
import pandas as pd
import pytesseract
from pdf2image import convert_from_bytes, convert_from_path
from PIL import Image
import re
import io
from pathlib import Path
import tempfile
import os
# Import deduplication utilities
from deduplication_utils import deduplicate_dataframe, compare_extraction_methods
# Import database utilities
from database_utils import save_to_database, get_database_stats, export_database_to_excel, get_database_path
import fitz  # PyMuPDF for efficient PDF handling
import logging
import traceback
from datetime import datetime
import sys
import subprocess
import zipfile
from typing import List, Dict, Any

# Set up logging (must be before any logger usage)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Word document processing
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

# Configure Tesseract OCR path. Use the hardcoded path only if tesseract is not in PATH.
try:
    # Check if Tesseract is accessible via PATH
    pytesseract.get_tesseract_version()
    logger.info("Tesseract found in system PATH.")
except pytesseract.TesseractNotFoundError:
    TESSERACT_PATH = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    if os.path.exists(TESSERACT_PATH):
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
        logger.warning(f"Tesseract not in PATH. Using hardcoded path: {TESSERACT_PATH}")
    else:
        # If the specific path also fails, log an error. OCR functions will fail gracefully.
        logger.error("Tesseract not found. Please ensure it is installed and in your PATH, or set the correct path in the code.")

# Configure page layout
st.set_page_config(
    page_title="KRA iTax - Data Extraction System",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Authentic KRA iTax styling based on official portal
st.markdown("""
<style>
    /* Import fonts similar to KRA */
    @import url('https://fonts.googleapis.com/css2?family=Arial:wght@400;500;600;700&display=swap');
    
    /* KRA Official Colors from iTax portal */
    :root {
        --kra-red: #dc2626;
        --kra-red-dark: #b91c1c;
        --kra-blue: #1e40af;
        --kra-blue-light: #3b82f6;
        --kra-gray: #f3f4f6;
        --kra-gray-dark: #6b7280;
        --kra-white: #ffffff;
        --kra-black: #1f2937;
    }
    
    /* Reset default Streamlit styling */
    .main .block-container {
        padding-top: 0rem;
        padding-bottom: 2rem;
        font-family: Arial, sans-serif;
        max-width: 100%;
    }
    
    /* KRA Header - Red bar like iTax */
    .kra-header-bar {
        background: var(--kra-red);
        color: white;
        padding: 0.5rem 2rem;
        margin: -1rem -1rem 0 -1rem;
        font-size: 0.9rem;
        text-align: center;
        font-weight: 500;
    }
    
    /* Main KRA Header */
    .kra-main-header {
        background: var(--kra-white);
        padding: 1.5rem 2rem;
        border-bottom: 2px solid var(--kra-gray);
        margin-bottom: 2rem;
    }
    
    .kra-logo-section {
        display: flex;
        align-items: center;
        justify-content: space-between;
        margin-bottom: 1rem;
    }
    
    .kra-title {
        font-size: 1.8rem;
        font-weight: 700;
        color: var(--kra-black);
        margin: 0;
    }
    
    .kra-subtitle {
        color: var(--kra-gray-dark);
        font-size: 1rem;
        margin: 0.5rem 0 0 0;
    }
    
    /* Service cards like iTax portal */
    .kra-services-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
        gap: 1.5rem;
        margin-bottom: 2rem;
    }
    
    .kra-service-card {
        background: var(--kra-white);
        border: 1px solid #e5e7eb;
        border-radius: 8px;
        padding: 1.5rem;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
        transition: box-shadow 0.3s ease;
    }
    
    .kra-service-card:hover {
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .kra-service-title {
        font-size: 1.1rem;
        font-weight: 600;
        color: var(--kra-blue);
        margin-bottom: 0.5rem;
    }
    
    .kra-service-desc {
        color: var(--kra-gray-dark);
        font-size: 0.9rem;
        line-height: 1.4;
    }
    
    /* Stats cards */
    .kra-stats-container {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 1rem;
        margin-bottom: 2rem;
    }
    
    .kra-stat-card {
        background: linear-gradient(135deg, var(--kra-blue) 0%, var(--kra-blue-light) 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 8px;
        text-align: center;
        box-shadow: 0 2px 4px rgba(30, 64, 175, 0.2);
    }
    
    .kra-stat-number {
        font-size: 2rem;
        font-weight: 700;
        margin-bottom: 0.25rem;
    }
    
    .kra-stat-label {
        font-size: 0.9rem;
        opacity: 0.9;
    }
    
    /* Buttons - KRA Red style */
    .stButton > button {
        background: var(--kra-red) !important;
        color: white !important;
        border: none !important;
        border-radius: 4px !important;
        padding: 0.75rem 1.5rem !important;
        font-weight: 500 !important;
        transition: background-color 0.3s ease !important;
    }
    
    .stButton > button:hover {
        background: var(--kra-red-dark) !important;
        color: white !important;
    }
    
    /* File uploader */
    .stFileUploader {
        border: 2px dashed var(--kra-blue);
        border-radius: 8px;
        padding: 2rem;
        background: #f8fafc;
        text-align: center;
    }
    
    /* Progress bars */
    .stProgress > div > div {
        background: var(--kra-red);
    }
    
    /* Tables */
    .stDataFrame {
        border: 1px solid #e5e7eb;
        border-radius: 8px;
        overflow: hidden;
    }
    
    /* Sidebar */
    .css-1d391kg, .css-1544g2n {
        background: var(--kra-gray);
        border-right: 1px solid #d1d5db;
    }
    
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stDeployButton {display: none;}
    header[data-testid="stHeader"] {display: none;}
    
    /* Success/Error styling */
    .stSuccess {
        background: #10b981;
        border-radius: 4px;
    }
    
    .stError {
        background: var(--kra-red);
        border-radius: 4px;
    }
    
    .stWarning {
        background: #f59e0b;
        border-radius: 4px;
    }
    
    .stInfo {
        background: var(--kra-blue);
        border-radius: 4px;
    }
</style>
""", unsafe_allow_html=True)

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Debug mode toggle
DEBUG_MODE = st.sidebar.toggle("🐛 Debug Mode", value=False, help="Enable detailed debugging information")

def log_debug(message, data=None):
    """Log debug information if debug mode is enabled"""
    if DEBUG_MODE:
        timestamp = datetime.now().strftime("%H:%M:%S")
        st.sidebar.write(f"🐛 **{timestamp}**: {message}")
        if data:
            st.sidebar.json(data)
        logger.info(f"DEBUG: {message}")

def log_error(message, error):
    """Log error information"""
    error_msg = f"ERROR: {message} - {str(error)}"
    st.error(error_msg)
    logger.error(error_msg)
    if DEBUG_MODE:
        st.sidebar.error(f"❌ {message}")
        st.sidebar.code(str(error))

def extract_text_from_word(file_path_or_bytes):
    """
    Extract text from Word document (.docx) including tables
    
    Args:
        file_path_or_bytes: File path string or bytes object
        
    Returns:
        tuple: (extracted_text, extraction_method)
    """
    try:
        log_debug("Starting Word document text extraction with table support")
        
        if isinstance(file_path_or_bytes, str):
            # File path
            if DOCX_AVAILABLE:
                doc = Document(file_path_or_bytes)
                
                # Extract paragraphs
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract tables
                for table in doc.tables:
                    table_text = "\n--- TABLE START ---\n"
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            table_text += " | ".join(row_text) + "\n"
                    table_text += "--- TABLE END ---\n"
                    text_parts.append(table_text)
                
                text = "\n".join(text_parts)
                if text.strip():
                    log_debug(f"Word extraction successful via python-docx with tables: {len(text)} characters")
                    return text, "docx_with_tables"
            
            if DOCX2TXT_AVAILABLE:
                text = docx2txt.process(file_path_or_bytes)
                if text.strip():
                    log_debug(f"Word extraction successful via docx2txt: {len(text)} characters")
                    return text, "docx2txt"
        else:
            # Bytes object (uploaded file)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(file_path_or_bytes)
                tmp_path = tmp_file.name
            
            try:
                if DOCX_AVAILABLE:
                    doc = Document(tmp_path)
                    
                    # Extract paragraphs
                    text_parts = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
                    
                    # Extract tables
                    for table in doc.tables:
                        table_text = "\n--- TABLE START ---\n"
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text:
                                    row_text.append(cell_text)
                            if row_text:
                                table_text += " | ".join(row_text) + "\n"
                        table_text += "--- TABLE END ---\n"
                        text_parts.append(table_text)
                    
                    text = "\n".join(text_parts)
                    if text.strip():
                        log_debug(f"Word extraction successful via python-docx with tables: {len(text)} characters")
                        return text, "docx_with_tables"
                
                if DOCX2TXT_AVAILABLE:
                    text = docx2txt.process(tmp_path)
                    if text.strip():
                        log_debug(f"Word extraction successful via docx2txt: {len(text)} characters")
                        return text, "docx2txt"
            finally:
                os.unlink(tmp_path)
        
        log_debug("Word text extraction failed - no text found")
        return "", "word_extraction_failed"
        
    except Exception as e:
        log_error("Error extracting text from Word document", e)
        return "", f"word_error: {str(e)}"

def extract_text_from_pdf(pdf_file):
    """
    Extract text from PDF - detect if digital or scanned, with enhanced table extraction
    
    Args:
        pdf_file: File path, bytes, or Streamlit UploadedFile object
        
    Returns:
        tuple: (extracted_text, extraction_method)
    """
    try:
        log_debug("Starting PDF text extraction with table support")
        
        # Handle different input types
        if hasattr(pdf_file, 'read'):
            # Streamlit UploadedFile
            pdf_bytes = pdf_file.read()
            pdf_file.seek(0)  # Reset for potential reuse
        elif isinstance(pdf_file, bytes):
            pdf_bytes = pdf_file
        else:
            # File path
            with open(pdf_file, 'rb') as f:
                pdf_bytes = f.read()
        
        # Try digital text extraction first with table detection
        try:
            pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            text_parts = []
            
            for page_num, page in enumerate(pdf_doc):
                page_text = page.get_text()
                
                # Try to extract tables using PyMuPDF's table detection
                try:
                    tables = page.find_tables()
                    if tables:
                        table_text = f"\n--- PAGE {page_num + 1} TABLES ---\n"
                        for table_num, table in enumerate(tables):
                            table_text += f"TABLE {table_num + 1}:\n"
                            table_data = table.extract()
                            for row in table_data:
                                if row and any(cell and str(cell).strip() for cell in row):
                                    row_text = " | ".join(str(cell).strip() if cell else "" for cell in row)
                                    table_text += row_text + "\n"
                            table_text += "\n"
                        text_parts.append(table_text)
                    
                    # Add regular text
                    if page_text.strip():
                        text_parts.append(page_text)
                        
                except Exception as table_error:
                    log_debug(f"Table extraction failed for page {page_num + 1}: {table_error}")
                    # Fall back to regular text extraction
                    if page_text.strip():
                        text_parts.append(page_text)
            
            pdf_doc.close()
            
            text = "\n".join(text_parts)
            if len(text.strip()) > 100:  # Meaningful text threshold
                log_debug(f"Digital PDF extraction with tables successful: {len(text)} characters")
                return text, "digital_pdf_with_tables"
        except Exception as e:
            log_debug(f"Digital PDF extraction failed: {e}")
        
        # Fall back to OCR with enhanced table detection
        try:
            images = convert_from_bytes(pdf_bytes, dpi=300)
            text_parts = []
            
            for i, image in enumerate(images):
                # Use OCR with table structure preservation
                custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
                page_text = pytesseract.image_to_string(image, config=custom_config)
                
                if page_text.strip():
                    text_parts.append(f"--- PAGE {i+1} (OCR) ---\n{page_text}")
                
                log_debug(f"OCR page {i+1}: {len(page_text)} characters")
            
            text = "\n".join(text_parts)
            if text.strip():
                log_debug(f"PDF OCR extraction with table support successful: {len(text)} characters")
                return text, "pdf_ocr_with_tables"
        except Exception as e:
            log_error("PDF OCR extraction failed", e)
        
        return "", "pdf_extraction_failed"
        
    except Exception as e:
        log_error("Error in PDF text extraction", e)
        return "", f"pdf_error: {str(e)}"

def extract_kra_fields(text):
    """
    Final improved KRA field extraction with only the specified 6 fields and comprehensive fixes
    
    Args:
        text: Extracted text from document
        
    Returns:
        dict: Dictionary containing extracted fields
    """
    data = {
        'date': '',
        'pin': '',
        'taxpayerName': '',
        'notice': '',
        'preAmount': '',
        'finalAmount': '',
        'year': '',
        'officerName': '',
        'station': ''
    }
    
    try:
        log_debug("Starting KRA field extraction")
        
        # Extract Date (after PIN, before taxpayer name/address)
        date_match = re.search(r'PIN[:\s]*[A-Z]\d{9}[A-Z][^\n\r]*[\n\r]+\s*([0-9]{1,2}(?:ST|ND|RD|TH)?\s+[A-Z]+,?\s+\d{4})', text, re.IGNORECASE)
        if not date_match:
            # Fallback: look for date after PIN, before a name/address (comma or P.O. BOX)
            date_match = re.search(r'PIN[:\s]*[A-Z]\d{9}[A-Z][^\n\r]*[\n\r]+\s*([0-9]{1,2}(?:ST|ND|RD|TH)?\s+[A-Z]+,?\s+\d{4})[\n\r]+[A-Z][A-Z\s,&.\-]+[,\n]', text, re.IGNORECASE)
        if date_match:
            data['date'] = date_match.group(1).strip()
        else:
            # Fallback to previous patterns
            date_patterns = [
                r'(\d{1,2}(?:ST|ND|RD|TH)\s+[A-Z]+,?\s+\d{4})',
                r'(\d{1,2}[A-Z]{2}\s+[A-Z]{3,9},?\s*\d{4})',
                r'(\d{1,2}\s+[A-Z]{3,9}\s+\d{4})',
                r'(\d{1,2}[-/]\d{1,2}[-/]\d{4})',
            ]
            for pattern in date_patterns:
                date_match = re.search(pattern, text, re.IGNORECASE)
                if date_match:
                    data['date'] = date_match.group(1).strip()
                    break
        
        # Extract PIN (existing patterns work well)
        pin_patterns = [
            r'PIN[:\s]*([A-Z]\d{9}[A-Z])',
            r'P\.?I\.?N\.?[:\s]*([A-Z]\d{9}[A-Z])',
            r'([A-Z]\d{9}[A-Z])',
        ]
        
        for pattern in pin_patterns:
            pin_match = re.search(pattern, text, re.IGNORECASE)
            if pin_match:
                pin = pin_match.group(1).upper()
                if re.match(r'^[A-Z]\d{9}[A-Z]$', pin):
                    data['pin'] = pin
                    break
        
        # IMPROVED: Extract Taxpayer Name (context-aware: after date, before address)
        # Look for name after date, before P.O. BOX or comma
        name_match = re.search(r'(?:\d{1,2}(?:ST|ND|RD|TH)?\s+[A-Z]+,?\s+\d{4})\s*\n\s*([A-Z][A-Z\s&.,\-]+?),', text, re.IGNORECASE)
        if not name_match:
            # Fallback: name after date, before P.O. BOX
            name_match = re.search(r'(?:\d{1,2}(?:ST|ND|RD|TH)?\s+[A-Z]+,?\s+\d{4})\s*\n\s*([A-Z][A-Z\s&.,\-]{5,}?)\s*\n\s*P\.\s*O\.', text, re.IGNORECASE)
        if not name_match:
            # Fallback: name after PIN, before address
            name_match = re.search(r'PIN[:\s]*[A-Z]\d{9}[A-Z][^\n\r]*[\n\r]+.*?([A-Z][A-Z\s&.,\-]{5,}?),', text, re.IGNORECASE | re.DOTALL)
        if name_match:
            name = name_match.group(1).strip()
            name = re.sub(r'\s+', ' ', name)
            name = name.replace('\n', ' ').replace('\r', ' ')
            name = re.sub(r'\s+', ' ', name)
            # Enhanced validation for both individual and business names
            words = name.split()
            valid_keywords = ['LIMITED', 'LTD', 'COMPANY', 'GROUP', 'CORPORATION', 'CORP', 'INC', 'ENTERPRISES', 'SERVICES']
            is_valid = (
                len(name) >= 5 and len(name) <= 100 and
                (
                    any(keyword in name.upper() for keyword in valid_keywords) or
                    (len(words) >= 2 and len(words) <= 4 and all(word.isalpha() for word in words))
                ) and
                sum(c.isalpha() or c.isspace() for c in name) / len(name) > 0.7
            )
            if is_valid:
                data['taxpayerName'] = name
        # Extract Station strictly as the line after the P.O. BOX address (no fallbacks)
        # Extract station ONLY as the line after the P.O. BOX address. No fallbacks.
        station_match = re.search(r'P\.\s*O\.\s*BOX[^\n\r]*[\n\r]+\s*([A-Z][A-Z\s.&,-]+)[\.,]', text)
        if station_match:
            station = station_match.group(1).strip()
            station = re.sub(r'\s+', ' ', station)
            station = station.rstrip('.').rstrip(',')
            data['station'] = station
        
        # NEW: Extract Notice (RE: content from letters)
        notice_patterns = [
            r'RE:\s*(.+?)(?:\n\n|\r\n\r\n|Dear|KRA|$)',  # RE: followed by content until double newline or certain words
            r'RE:\s*(.+?)(?:\n[A-Z]|\r\n[A-Z]|$)',      # RE: followed by content until next line starting with capital
            r'RE:\s*(.+?)(?=\n\s*\n|\r\n\s*\r\n|$)',    # RE: followed by content until empty line
            r'RE:\s*(.+?)\.?\s*(?:\n|$)',               # RE: followed by content until end of line
        ]
        
        for pattern in notice_patterns:
            notice_match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if notice_match:
                notice = notice_match.group(1).strip()
                # Clean up the notice text
                notice = re.sub(r'\s+', ' ', notice)  # Replace multiple spaces/newlines with single space
                notice = notice.replace('\n', ' ').replace('\r', ' ')  # Remove line breaks
                notice = re.sub(r'\s+', ' ', notice).strip()  # Final cleanup
                
                # Validate notice (should be meaningful text)
                if len(notice) >= 10 and len(notice) <= 200:  # Reasonable length
                    data['notice'] = notice
                    break
        
        # NEW: Extract Pre-Amount (Total Tax) - ENHANCED with table extraction
        # Step 1: Look for table structure markers and extract amounts from them
        table_amount_patterns = [
            # Table row patterns with separators
            r'total[\s\|]*tax[\s\|]*([0-9,]+\.?\d*)',
            r'tax[\s\|]*amount[\s\|]*([0-9,]+\.?\d*)',
            r'amount[\s\|]*due[\s\|]*([0-9,]+\.?\d*)',
            # Patterns within TABLE START/END markers
            r'---\s*table\s+start\s*---.*?total.*?([0-9,]+\.?\d*).*?---\s*table\s+end\s*---',
            # Pipe-separated table patterns
            r'\|[^|]*total[^|]*\|[^|]*([0-9,]+\.?\d*)[^|]*\|',
            r'\|[^|]*tax[^|]*\|[^|]*([0-9,]+\.?\d*)[^|]*\|',
        ]
        
        for pattern in table_amount_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                amount = match.group(1).strip()
                clean_amount = amount.replace(',', '')
                try:
                    amount_value = float(clean_amount)
                    if amount_value >= 10:  # Minimum threshold
                        data['preAmount'] = amount
                        break
                except ValueError:
                    continue
        
        # Step 2: Look for lines containing "Total Tax" specifically (existing logic)
        if not data['preAmount']:
            total_tax_lines = []
            for line in text.split('\n'):
                if re.search(r'total\s+tax', line, re.IGNORECASE):
                    total_tax_lines.append(line.strip())
            
            # Step 3: Extract amount from Total Tax lines only
            if total_tax_lines:
                for line in total_tax_lines:
                    # Pattern to find amount on the same line as "Total Tax"
                    # Matches formats like: "Total Tax 14,769.50", "Total Tax: 14,769.50", "Total Tax                14,769.50"
                    amount_patterns = [
                        r'total\s+tax[:\s]*([0-9,]+\.?\d*)',  # Total Tax followed by amount with commas
                        r'total\s+tax.*?([0-9,]+\.?\d*)',     # Total Tax with anything in between, then amount with commas
                        r'total\s+tax[:\s]*([0-9]+\.?\d*)',   # Total Tax followed by amount WITHOUT commas
                        r'total\s+tax.*?([0-9]+\.?\d*)',      # Total Tax with anything in between, then amount WITHOUT commas
                        r'([0-9,]+\.?\d*)\s*total\s+tax',     # Amount before Total Tax (reversed order)
                        r'([0-9]+\.?\d*)\s*total\s+tax',      # Amount WITHOUT commas before Total Tax
                    ]
                    
                    for pattern in amount_patterns:
                        match = re.search(pattern, line, re.IGNORECASE)
                        if match:
                            amount = match.group(1).strip()
                            # Validate the amount
                            clean_amount = amount.replace(',', '')
                            try:
                                amount_value = float(clean_amount)
                                if amount_value >= 10:  # Minimum threshold
                                    data['preAmount'] = amount  # Keep original formatting with commas
                                    break
                            except ValueError:
                                continue
                    
                    # If we found amount on this Total Tax line, stop searching
                    if data['preAmount']:
                        break
        
        # Step 3: Fallback - if no "Total Tax" line found, try other tax amount patterns
        if not data['preAmount']:
            fallback_patterns = [
                r'(?:principal\s+tax|tax\s+due|final\s+tax)[:\s]*([0-9,]+\.?\d*)',
                r'(?:amount\s+due|net\s+tax|payable)[:\s]*([0-9,]+\.?\d*)',
                # More aggressive patterns for tax amounts
                r'(?:tax\s+amount|amount\s+of\s+tax|tax\s+payable)[:\s]*([0-9,]+\.?\d*)',
                r'(?:withholding\s+tax|paye\s+tax|income\s+tax)[:\s]*([0-9,]+\.?\d*)',
                # Look for amounts near the word "tax" (broader search)
                r'tax[^\d]*?([0-9,]+\.?\d*)',
                # Look for any substantial amount (as last resort for documents with tax data)
                r'\b([0-9]{1,2}(?:,\d{3})+\.\d{2})\b',  # Pattern like 14,769.50
            ]
            for pattern in fallback_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    amount = match.group(1).strip()
                    clean_amount = amount.replace(',', '')
                    # Skip if the amount looks like a PIN (10-11 alphanumeric chars, e.g. A123456789X)
                    if re.fullmatch(r'[A-Z]\d{9}[A-Z]', amount, re.IGNORECASE):
                        continue
                    if len(clean_amount) == 11 and clean_amount.isdigit():
                        continue
                    try:
                        amount_value = float(clean_amount)
                        # More lenient threshold for fallback patterns
                        if amount_value >= 100:  # Higher minimum for fallback to avoid false positives
                            data['preAmount'] = amount
                            break
                    except ValueError:
                        continue
                if data['preAmount']:
                    break
        
        # FIXED: Extract Year with proper patterns and business logic
        year_found = False
        
        # Try ONLY explicit tax year mentions (no P.O. BOX ranges)
        explicit_year_patterns = [
            r'(?:tax\s+year|year\s+of\s+income|for\s+the\s+year)[:\s]*(\d{4})',
            # Only year ranges in tax contexts (not P.O. BOX)
            r'(?:tax\s+year|income\s+year|assessment).*?(\d{4}[-–]\d{4})',
        ]
        
        for pattern in explicit_year_patterns:
            year_match = re.search(pattern, text, re.IGNORECASE)
            if year_match:
                year = year_match.group(1).strip()
                if year.isdigit() and 2015 <= int(year) <= 2030:
                    data['year'] = year
                    year_found = True
                    break
                elif '-' in year and len(year.split('-')[0]) == 4:  # Valid year range
                    data['year'] = year
                    year_found = True
                    break
        
        # If no explicit year found, use business logic: document year - 1
        if not year_found and data['date']:
            # Extract year from document date
            doc_year_match = re.search(r'\d{4}', data['date'])
            if doc_year_match:
                doc_year = int(doc_year_match.group(0))
                # Tax assessments are typically for the previous year
                tax_year = doc_year - 1
                data['year'] = str(tax_year)
                year_found = True
        
        # IMPROVED: Extract Officer Name (from contact information)
        improved_officer_patterns = [
            # Pattern 1: Direct Officer mention (most direct)
            r'Officer[:\s]*([A-Z][a-zA-Z\s]+?)(?:\n|Contact|Tel|Phone|Email)',
            # Pattern 2: Contact name in "contact X or Y" phrase (most reliable for this doc type)
            r'contact\s+([A-Z][a-z]+\s+[A-Z][a-z]+)\s+or',
            # Pattern 3: After "hesitate to contact"
            r'hesitate\s+to\s+contact\s+([A-Z][a-z]+\s+[A-Z][a-z]+)',
            # Pattern 4: After "contact" before phone number
            r'contact\s+([A-Z][a-z]+\s+[A-Z][a-z]+).*?phone',
            # Pattern 5: Signature name (fallback)
            r'Yours\s+faithfully,.*?\n\s*([A-Z][a-z]+\s+[A-Z][a-z]+)',
            # Pattern 6: General contact pattern
            r'contact\s+([A-Z][a-z]+\s+[A-Z][a-z]+)',
        ]
        
        for pattern in improved_officer_patterns:
            officer_match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if officer_match:
                officer = officer_match.group(1).strip()
                officer = re.sub(r'\s+', ' ', officer)
                
                # Enhanced validation - should be 2-4 words, all letters/spaces
                words = officer.split()
                if (len(words) >= 2 and len(words) <= 4 and 
                    all(word.isalpha() for word in words) and 
                    len(officer) >= 5 and len(officer) <= 50):
                    data['officerName'] = officer
                    break
        
        # Extract Station (existing patterns work well)
        station_patterns = [
            # Priority 1: Station mentioned with "STATION" or "OFFICE" (most specific)
            r'([A-Z]{4,})\s+(?:STATION|OFFICE)',
            # Priority 2: Known KRA stations in document (avoid KRA headquarters in header)
            r'(?:P\.?\s*O\.?\s*BOX\s+\d+[^,\n]*,?\s*)?([A-Z]{4,}(?:LODWAR|MOMBASA|KISUMU|NAKURU|ELDORET|NYERI|MERU|MACHAKOS|KITALE|GARISSA|ISIOLO|MALINDI|KILIFI|EMBU|THIKA|KIAMBU|KAKAMEGA|KERICHO|BOMET|BUNGOMA|WEBUYE|MIGORI|HOMABAY|SIAYA|BUSIA|MARSABIT|MANDERA|WAJIR|MOYALE|KAPENGURIA|MARALAL))',
            # Priority 3: Specific station names (excluding NAIROBI from KRA header)
            r'\b(LODWAR|MOMBASA|KISUMU|NAKURU|ELDORET|NYERI|MERU|MACHAKOS|KITALE|GARISSA|ISIOLO|MALINDI|KILIFI|EMBU|THIKA|KIAMBU|KAKAMEGA|KERICHO|BOMET|BUNGOMA|WEBUYE|MIGORI|HOMABAY|SIAYA|BUSIA|MARSABIT|MANDERA|WAJIR|MOYALE|KAPENGURIA|MARALAL)\b',
            # Priority 4: General location after P.O. BOX (fallback)
            r'P\.?\s*O\.?\s*BOX\s+\d+[-–\s]*\d*[,\s]*([A-Z]{3,})',
        ]
        
        for pattern in station_patterns:
            station_match = re.search(pattern, text, re.IGNORECASE)
            if station_match:
                station = station_match.group(1).strip().upper()
                if len(station) >= 3:
                    data['station'] = station
                    break
        
        fields_found = sum(1 for v in data.values() if v)
        log_debug(f"KRA field extraction completed: {fields_found}/6 fields found")
        
        return data
        
    except Exception as e:
        log_error("Error in KRA field extraction", e)
        return data

def process_document(file_path_or_uploaded, file_name):
    """
    Process a single document (PDF or Word) and extract KRA data
    
    Args:
        file_path_or_uploaded: File path or uploaded file object
        file_name: Name of the file for identification
        
    Returns:
        dict: Processing results with extracted data
    """
    # Initialize result with only the 8 core fields (camelCase)
    result = {field: '' for field in ['date', 'pin', 'taxpayerName', 'notice', 'preAmount', 'finalAmount', 'year', 'officerName', 'station']}
    
    try:
        log_debug(f"Processing document: {file_name}")
        # Determine file type
        file_ext = Path(file_name).suffix.lower()
        # Extract text based on file type
        if file_ext == '.pdf':
            text, method = extract_text_from_pdf(file_path_or_uploaded)
        elif file_ext in ['.docx', '.doc']:
            if hasattr(file_path_or_uploaded, 'read'):
                text, method = extract_text_from_word(file_path_or_uploaded.read())
            else:
                text, method = extract_text_from_word(file_path_or_uploaded)
        else:
            log_debug(f"Unsupported file type: {file_ext}")
            return result
        if not text:
            log_debug("No text extracted from document")
            return result
        # Extract KRA fields
        kra_data = extract_kra_fields(text)
        result.update(kra_data)
        # Store extracted text for debugging (first file only)
        if not hasattr(st.session_state, 'last_extracted_text'):
            st.session_state.last_extracted_text = text
        # Store all extracted raw texts for Data Portal access
        if 'raw_texts' not in st.session_state:
            st.session_state.raw_texts = {}
        st.session_state.raw_texts[file_name] = text
        # --- DEBUG OUTPUT: Show extracted text and preAmount lines ---
        if 'debug_outputs' not in st.session_state:
            st.session_state.debug_outputs = []
        debug_info = {
            'file': file_name,
            'type': method if 'method' in locals() else 'unknown',
            'text_preview': text[:500],
            'preAmount': result.get('preAmount', ''),
            'preAmount_lines': []
        }
        # Find lines with 'total tax' or numbers that could be matched
        lines = text.split('\n')
        for line in lines:
            if 'total tax' in line.lower() or (result.get('preAmount') and result['preAmount'] in line):
                debug_info['preAmount_lines'].append(line.strip())
        st.session_state.debug_outputs.append(debug_info)
        log_debug(f"Document processed successfully")
    except Exception as e:
        log_error(f"Error processing document {file_name}", e)
    return result

def process_folder(folder_path):
    """
    Process all supported documents in a folder
    
    Args:
        folder_path: Path to folder containing documents
        
    Returns:
        list: List of processing results
    """
    results = []
    supported_extensions = ['.pdf', '.docx', '.doc']
    
    try:
        folder = Path(folder_path)
        if not folder.exists():
            st.error(f"Folder does not exist: {folder_path}")
            return results
        
        # Find all supported files
        files = []
        for ext in supported_extensions:
            files.extend(folder.glob(f"*{ext}"))
            files.extend(folder.glob(f"*{ext.upper()}"))
        
        if not files:
            st.warning(f"No supported documents found in folder. Looking for: {', '.join(supported_extensions)}")
            return results
        
        st.info(f"Found {len(files)} documents to process")
        
        # Process each file
        progress_bar = st.progress(0)
        status_placeholder = st.empty()
        
        for i, file_path in enumerate(files):
            progress = (i + 1) / len(files)
            progress_bar.progress(progress)
            status_placeholder.info(f"Processing {i+1}/{len(files)}: {file_path.name}")
            
            result = process_document(str(file_path), file_path.name)
            results.append(result)
        
        progress_bar.progress(1.0)
        status_placeholder.success(f"Completed processing {len(files)} documents")
        
    except Exception as e:
        log_error("Error processing folder", e)
    
    return results

def main():
    # Enhanced KRA iTax Professional UI Styling
    st.markdown("""
    <style>
    /* Import professional fonts */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Roboto:wght@300;400;500;700&display=swap');
    
    /* Global font and layout improvements */
    .stApp {
        font-family: 'Inter', 'Roboto', 'Segoe UI', sans-serif !important;
    }
    
    /* Remove default margins and improve spacing */
    .main .block-container {
        padding-top: 1rem !important;
        padding-bottom: 1rem !important;
        max-width: 1200px !important;
    }
    
    /* Navigation Card - Clean white design */
    .nav-container {
        background: white;
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 2rem;
        box-shadow: 0 2px 10px rgba(0,0,0,0.08);
        border: 1px solid #E5E7EB;
        position: sticky;
        top: 10px;
        z-index: 100;
    }
    
    .nav-title {
        color: #374151 !important;
        font-size: 1.3rem !important;
        font-weight: 600 !important;
        margin-bottom: 1rem !important;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }
    
    /* Radio button improvements */
    .stRadio > div {
        background: #FAFAFA;
        border-radius: 10px;
        padding: 1rem;
        border: 1px solid #E5E7EB;
    }
    
    .stRadio label {
        font-weight: 500 !important;
        color: #374151 !important;
        padding: 0.5rem 1rem !important;
        border-radius: 8px !important;
        transition: all 0.2s ease !important;
    }
    
    .stRadio label:hover {
        background: #F3F4F6 !important;
        color: #E31E24 !important;
    }
    
    /* Enhanced Metric Cards */
    div[data-testid="metric-container"] {
        background: white !important;
        border: 1px solid #E5E7EB !important;
        border-radius: 10px !important;
        padding: 1.2rem !important;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1) !important;
        transition: all 0.2s ease !important;
    }
    
    div[data-testid="metric-container"]:hover {
        box-shadow: 0 4px 12px rgba(227,30,36,0.15) !important;
        border-color: #E31E24 !important;
        transform: translateY(-2px) !important;
    }
    
    div[data-testid="metric-container"] label {
        font-size: 0.9rem !important;
        font-weight: 500 !important;
        color: #6B7280 !important;
    }
    
    div[data-testid="metric-container"] div[data-testid="metric-value"] {
        font-size: 1.8rem !important;
        font-weight: 700 !important;
        color: #111827 !important;
    }
    
    /* Section Cards */
    .section-card {
        background: white;
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1.5rem;
        box-shadow: 0 2px 10px rgba(0,0,0,0.08);
        border: 1px solid #F3F4F6;
    }
    
    /* Enhanced Buttons */
    .stButton > button {
        background: linear-gradient(135deg, #E31E24 0%, #B91C1C 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        padding: 0.6rem 1.5rem !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 2px 4px rgba(227,30,36,0.2) !important;
    }
    
    .stButton > button:hover {
        background: linear-gradient(135deg, #B91C1C 0%, #991B1B 100%) !important;
        box-shadow: 0 4px 12px rgba(227,30,36,0.4) !important;
        transform: translateY(-1px) !important;
    }
    
    /* Secondary Buttons */
    .stButton > button[kind="secondary"] {
        background: white !important;
        color: #E31E24 !important;
        border: 2px solid #E31E24 !important;
    }
    
    .stButton > button[kind="secondary"]:hover {
        background: #FEF2F2 !important;
        border-color: #B91C1C !important;
        color: #B91C1C !important;
    }
    
    /* Download Buttons */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #10B981 0%, #059669 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        padding: 0.6rem 1.5rem !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 2px 4px rgba(16,185,129,0.2) !important;
    }
    
    .stDownloadButton > button:hover {
        background: linear-gradient(135deg, #059669 0%, #047857 100%) !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 4px 12px rgba(16,185,129,0.4) !important;
    }
    
    /* Headers and Typography */
    h1, h2, h3 {
        font-family: 'Inter', sans-serif !important;
        font-weight: 600 !important;
    }
    
    h2 {
        color: #E31E24 !important;
        font-size: 1.8rem !important;
        margin-bottom: 1rem !important;
        padding-bottom: 0.5rem !important;
        border-bottom: 2px solid #E31E24 !important;
    }
    
    h3 {
        color: #374151 !important;
        font-size: 1.3rem !important;
        margin-bottom: 0.8rem !important;
    }
    
    /* File Uploader Enhancement */
    .stFileUploader > div > div {
        border: 2px dashed #E31E24 !important;
        border-radius: 12px !important;
        background: #FEFEFE !important;
        padding: 2rem !important;
        transition: all 0.3s ease !important;
    }
    
    .stFileUploader > div > div:hover {
        border-color: #B91C1C !important;
        background: #FEF2F2 !important;
    }
    
    /* Progress Bar */
    .stProgress > div > div {
        background: linear-gradient(90deg, #E31E24 0%, #B91C1C 100%) !important;
        border-radius: 10px !important;
    }
    
    /* Data Tables */
    .stDataFrame {
        border: 1px solid #E5E7EB !important;
        border-radius: 10px !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06) !important;
    }
    
    /* Alert Messages */
    .stSuccess {
        background: linear-gradient(135deg, #10B981 0%, #059669 100%) !important;
        color: white !important;
        border-radius: 8px !important;
        border: none !important;
        font-weight: 500 !important;
    }
    
    .stInfo {
        background: linear-gradient(135deg, #3B82F6 0%, #1D4ED8 100%) !important;
        color: white !important;
        border-radius: 8px !important;
        border: none !important;
        font-weight: 500 !important;
    }
    
    .stWarning {
        background: linear-gradient(135deg, #F59E0B 0%, #D97706 100%) !important;
        color: white !important;
        border-radius: 8px !important;
        border: none !important;
        font-weight: 500 !important;
    }
    
    /* Hide Streamlit elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Sidebar completely hidden */
    .css-1d391kg {display: none !important;}
    .css-1x8cf1d {display: none !important;}
    
    </style>
    """, unsafe_allow_html=True)
    
    # Navigation Section - Fixed at top
    st.markdown("""
    <div class="nav-container">
        <div class="nav-title">📋 Select Processing Mode</div>
    </div>
    """, unsafe_allow_html=True)
    
    processing_mode = st.radio(
        "",  # Empty label since we have custom title above
        ["📄 Upload Individual Files", "📁 Process Folder Batch"],
        index=0,
        help="Choose between uploading individual files or processing all documents in a folder",
        horizontal=True
    )
    
    # Main Content Area
    if processing_mode == "📄 Upload Individual Files":
        # File Upload Section
        st.markdown("""
        <div class="section-card">
        """, unsafe_allow_html=True)
        
        st.header("📄 Individual File Processing")
        
        # Database Information with refresh capability
        st.subheader("📊 Database Status")
        
        # Add refresh button
        col_refresh1, col_refresh2 = st.columns([3, 1])
        with col_refresh2:
            if st.button("🔄 Refresh", help="Refresh database statistics", key="refresh_db_stats"):
                st.rerun()
        
        # Get database stats (with cache invalidation if refresh triggered)
        db_stats = get_database_stats()
        
        # Check if database file exists
        db_path = get_database_path()
        db_exists = os.path.exists(db_path)
        
        if db_exists:
            st.success(f"📁 Database file: `{os.path.basename(db_path)}` ({'exists' if db_exists else 'not found'})")
        else:
            st.info("📁 Database file will be created after first extraction")
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Records", db_stats['total_records'])
        with col2:
            st.metric("Unique Taxpayers", db_stats['unique_taxpayers'])
        with col3:
            st.metric("Unique Stations", db_stats['unique_stations'])
        with col4:
            if db_stats['total_records'] > 0:
                # Add full database download button
                excel_data = export_database_to_excel()
                if excel_data:
                    st.download_button(
                        label="📥 Download Full Database",
                        data=excel_data,
                        file_name=f"KRA_Complete_Database_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        type="primary",
                        help="Download complete database with all historical records"
                    )
        
        if db_stats['total_records'] > 0:
            st.info(f"📅 Last updated: {db_stats['last_updated']} | 📊 Date range: {db_stats['date_range']}")
        
        st.subheader("📄 Upload Documents")
        
        # Initialize session state for file management
        if 'processed_files' not in st.session_state:
            st.session_state.processed_files = False
        if 'processing_results' not in st.session_state:
            st.session_state.processing_results = None
        
        # File uploader with key to enable clearing
        uploaded_files = st.file_uploader(
            "Upload documents (PDF or Word)",
            type=['pdf', 'docx', 'doc'],
            accept_multiple_files=True,
            help="Upload one or more PDF or Word documents containing KRA tax notices",
            key="file_uploader_main"
        )
        
        if uploaded_files and not st.session_state.processed_files:
            st.success(f"📁 {len(uploaded_files)} file(s) uploaded")
            
            if st.button("🚀 Process Uploaded Files", type="primary", key="process_button_main"):
                results = []
                
                progress_bar = st.progress(0)
                status_placeholder = st.empty()
                
                for i, uploaded_file in enumerate(uploaded_files):
                    progress = (i + 1) / len(uploaded_files)
                    progress_bar.progress(progress)
                    status_placeholder.info(f"Processing {i+1}/{len(uploaded_files)}: {uploaded_file.name}")
                    
                    # Process the uploaded file
                    result = process_document(uploaded_file, uploaded_file.name)
                    results.append(result)
                
                # Clear progress indicators
                progress_bar.empty()
                status_placeholder.empty()
                
                # Store results in session state
                st.session_state.processing_results = results
                st.session_state.processed_files = True
                
                # Force database stats refresh for updated display
                if 'db_stats_refresh' not in st.session_state:
                    st.session_state.db_stats_refresh = 0
                st.session_state.db_stats_refresh += 1
                
                st.success(f"✅ Successfully processed {len(uploaded_files)} file(s)!")
        
        # Display results if files have been processed
        if st.session_state.processed_files and st.session_state.processing_results:
            display_results(st.session_state.processing_results)
            
            # Add option to clear results and process new files
            if st.button("🔄 Process New Files", type="secondary", key="clear_results_button"):
                st.session_state.processed_files = False
                st.session_state.processing_results = None
                st.rerun()
        
        # Close section card
        st.markdown("</div>", unsafe_allow_html=True)
    
    else:  # Folder processing
        # Folder Processing Section
        st.markdown("""
        <div class="section-card">
        """, unsafe_allow_html=True)
        
        st.header("📁 Folder Batch Processing")
        
        # Database Information for folder processing too
        st.subheader("📊 Database Status")
        
        # Get database stats
        db_stats = get_database_stats()
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Records", db_stats['total_records'])
        with col2:
            st.metric("Unique Taxpayers", db_stats['unique_taxpayers'])
        with col3:
            st.metric("Unique Stations", db_stats['unique_stations'])
        with col4:
            if db_stats['total_records'] > 0:
                # Add full database download button
                excel_data = export_database_to_excel()
                if excel_data:
                    st.download_button(
                        label="📥 Download Full Database",
                        data=excel_data,
                        file_name=f"KRA_Complete_Database_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        type="primary",
                        help="Download complete database with all historical records",
                        key="folder_download_db"
                    )
        
        if db_stats['total_records'] > 0:
            st.info(f"📅 Last updated: {db_stats['last_updated']} | 📊 Date range: {db_stats['date_range']}")
        
        # Initialize session state for folder processing
        if 'folder_processed' not in st.session_state:
            st.session_state.folder_processed = False
        if 'folder_results' not in st.session_state:
            st.session_state.folder_results = None
        if 'folder_path_processed' not in st.session_state:
            st.session_state.folder_path_processed = ""
        
        folder_path = st.text_input(
            "Enter folder path containing documents:",
            placeholder="C:\\path\\to\\your\\documents",
            help="Enter the full path to the folder containing PDF and Word documents. Use forward slashes (/) or double backslashes (\\\\)",
            key="folder_path_input"
        )
        
        # Normalize the path for Windows compatibility
        if folder_path:
            folder_path = folder_path.strip().replace('/', '\\')
            if not folder_path.endswith('\\'):
                folder_path = folder_path
            st.info(f"📁 Looking for documents in: `{folder_path}`")
        
        if folder_path and not st.session_state.folder_processed:
            if st.button("🚀 Process Folder", type="primary", key="process_folder_button"):
                # Validate path before processing
                if not os.path.exists(folder_path):
                    st.error(f"❌ Folder does not exist: `{folder_path}`")
                    st.info("💡 **Tips:**\n- Use full path like `C:\\Users\\Username\\Documents\\KRA_Files`\n- Use forward slashes: `C:/Users/Username/Documents/KRA_Files`\n- Or double backslashes: `C:\\\\Users\\\\Username\\\\Documents\\\\KRA_Files`")
                else:
                    with st.spinner(f"🔍 Processing documents in {folder_path}..."):
                        results = process_folder(folder_path)
                        if results:
                            st.session_state.folder_results = results
                            st.session_state.folder_processed = True
                            st.session_state.folder_path_processed = folder_path
                    st.rerun()
        
        # Display folder results if processing is complete
        if st.session_state.folder_processed and st.session_state.folder_results:
            st.success(f"✅ Processed folder: {st.session_state.folder_path_processed}")
            display_results(st.session_state.folder_results)
            
            # Add button to reset for new folder
            if st.button("🔄 Process New Folder", type="secondary", key="reset_folder_button"):
                st.session_state.folder_processed = False
                st.session_state.folder_results = None
                st.session_state.folder_path_processed = ""
                st.rerun()
        
        # Close section card
        st.markdown("</div>", unsafe_allow_html=True)

def display_results(results):
    """Display processing results and save to database"""
    if not results:
        st.warning("No results to display")
        return
    
    st.header("📊 Extraction Results")
    
    # Create DataFrame from current results
    current_df = pd.DataFrame(results)
    
    # Apply deduplication to current batch
    deduplicated_current = deduplicate_dataframe(current_df)
    
    if len(deduplicated_current) < len(current_df):
        st.info(f"🔍 Removed {len(current_df) - len(deduplicated_current)} duplicate(s) from current batch")
    
    # Save to database automatically
    total_records, new_records, duplicates_removed = save_to_database(deduplicated_current, "multi_format_extractor")
    
    # Display save results
    col1, col2, col3 = st.columns(3)
    with col1:
        st.success(f"✅ {new_records} new record(s) added")
    with col2:
        st.info(f"📊 Total database records: {total_records}")
    with col3:
        if duplicates_removed > 0:
            st.warning(f"🔍 {duplicates_removed} duplicate(s) found and merged")
        else:
            st.success("🎉 No duplicates found")
    
    # Display current batch results
    st.subheader("📋 Current Batch Results")
    
    # DEBUG: Show extraction details for troubleshooting
    with st.expander("🔍 Debug: Extraction Details", expanded=False):
        debug_outputs = getattr(st.session_state, 'debug_outputs', [])
        for i, debug in enumerate(debug_outputs[:3]):
            st.write(f"**File {i+1} Debug Info:**")
            st.write(f"- File: `{debug.get('file', 'N/A')}`")
            st.write(f"- Type: `{debug.get('type', 'N/A')}`")
            st.write(f"- PIN: `{results[i].get('pin', 'NOT FOUND')}`")
            st.write(f"- Taxpayer Name: `{results[i].get('taxpayerName', 'NOT FOUND')}`")
            st.write(f"- Notice (RE:): `{results[i].get('notice', 'NOT FOUND')}`")
            st.write(f"- **preAmount: `{results[i].get('preAmount', 'NOT FOUND')}`**")
            st.write(f"- Year: `{results[i].get('year', 'NOT FOUND')}`")
            st.write(f"- Officer: `{results[i].get('officerName', 'NOT FOUND')}`")
            st.write("**Raw Extracted Text (first 500 chars):**")
            st.text(debug.get('text_preview', 'No text found'))
            if debug.get('preAmount_lines'):
                st.write("**preAmount-matched lines:**")
                for line in debug['preAmount_lines']:
                    st.code(line)
            else:
                st.write("❌ **No preAmount-matched lines found in document**")
            st.write("---")
    
    # Show the data in a nice table
    st.dataframe(
        deduplicated_current,
        use_container_width=True,
        hide_index=True
    )
    
    # Summary statistics for current batch
    st.subheader("📈 Current Batch Summary")
    col1, col2, col3, col4 = st.columns(4)
    
    total_files = len(results)
    successful = len([r for r in results if any(r.get(field, '') for field in ['date', 'pin', 'taxpayerName', 'notice', 'preAmount', 'finalAmount', 'year', 'officerName', 'station'])])
    success_rate = (successful / total_files * 100) if total_files > 0 else 0
    
    with col1:
        st.metric("Files Processed", total_files)
    with col2:
        st.metric("Successful Extractions", successful)
    with col3:
        st.metric("Success Rate", f"{success_rate:.1f}%")
    with col4:
        st.metric("Records Added to DB", new_records)
    
    # Download options
    st.subheader("📥 Download Options")
    col1, col2 = st.columns(2)
    
    with col1:
        # Current batch download
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            deduplicated_current.to_excel(writer, sheet_name='Current_Batch_Results', index=False)
        
        st.download_button(
            label="📥 Download Current Batch",
            data=output.getvalue(),
            file_name=f"KRA_Batch_Results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            help="Download results from this processing session"
        )
    
    with col2:
        # Full database download
        excel_data = export_database_to_excel()
        if excel_data:
            st.download_button(
                label="📥 Download Complete Database",
                data=excel_data,
                file_name=f"KRA_Complete_Database_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                help="Download entire database with all historical records"
            )
    
    # Force refresh of database stats in session state after save
    if 'db_stats_refresh' not in st.session_state:
        st.session_state.db_stats_refresh = 0
    st.session_state.db_stats_refresh += 1

if __name__ == "__main__":
    main()