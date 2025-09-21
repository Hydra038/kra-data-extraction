"""
Test Word Document Processing for KRA Extractor
===============================================

This script creates a test Word document with sample KRA data
to verify that the multi-format extractor can handle Word documents properly.
"""

try:
    from docx import Document
    from datetime import datetime
    import os
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('KENYA REVENUE AUTHORITY', 0)
    title.alignment = 1  # Center alignment
    
    # Add a subtitle
    subtitle = doc.add_heading('NOTICE UNDER SECTION 37 OF THE TAX PROCEDURES ACT', level=1)
    subtitle.alignment = 1
    
    # Add date
    doc.add_paragraph(f'26TH AUGUST, 2025')
    
    # Add taxpayer details
    doc.add_paragraph('')
    doc.add_paragraph('PIN: A123456789B')
    doc.add_paragraph('TEST COMPANY LIMITED')
    doc.add_paragraph('P.O. BOX 12345-00100')
    doc.add_paragraph('NAIROBI')
    
    # Add notice content
    doc.add_paragraph('')
    doc.add_paragraph('RE: NOTICE TO PAY OUTSTANDING TAX LIABILITY')
    doc.add_paragraph('')
    doc.add_paragraph('This notice is to inform you of outstanding tax liability for the year 2024.')
    doc.add_paragraph('')
    doc.add_paragraph('Details of tax liability:')
    doc.add_paragraph('Year of Income: 2024')
    doc.add_paragraph('Total Tax: 150,000.00')
    doc.add_paragraph('')
    doc.add_paragraph('Please settle this amount within 30 days from the date of this notice.')
    doc.add_paragraph('')
    doc.add_paragraph('Yours faithfully,')
    doc.add_paragraph('')
    doc.add_paragraph('JOHN MWANGI')
    doc.add_paragraph('Senior Tax Officer')
    doc.add_paragraph('NAIROBI STATION')
    
    # Save the document
    filename = 'test_kra_document.docx'
    doc.save(filename)
    
    print(f"✅ Test Word document created successfully: {filename}")
    print(f"📄 File size: {os.path.getsize(filename)} bytes")
    
except ImportError:
    print("❌ python-docx not available. Cannot create test document.")
except Exception as e:
    print(f"❌ Error creating test document: {e}")